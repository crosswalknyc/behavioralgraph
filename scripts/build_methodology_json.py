#!/usr/bin/env python3
"""
Convert the Crosswalk Product Methodologies .docx into the JSON schema
consumed by the dashboard's methodology modal
(/static/cw_methodologies.json).

Schema per section:
  { "title": str, "intro": [Paragraph, ...], "subs": [Sub, ...] }
Where:
  Paragraph = { "text": str, "is_list": bool, "runs": [{ "text": str, "bold": bool }, ...] }
  Sub       = { "title": str, "paragraphs": [Paragraph, ...] }

Sections are triggered by:
  - Heading 1 styled paragraphs
  - Normal paragraphs whose only run is bold and 17pt (215900 EMU) - matches
    'Glossary' and 'Contacts' in the July 28 doc, which are visually H1
    but not styled as H1 in Word.

Subsections are triggered by Heading 2 styled paragraphs.

The preamble (everything before the first section trigger) becomes a
_preamble_ section.
"""
import json
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

HEADING1_FONT_SIZE_EMU = 215900  # 17pt - matches both Heading 1 and Glossary/Contacts


def _iter_runs_including_hyperlinks(paragraph):
    """Yield (text, bold) for every visible run in a paragraph.

    python-docx's Paragraph.runs SKIPS runs nested inside <w:hyperlink>
    elements, so email links in Contacts / any linked text get silently
    dropped. Walk the XML directly to catch both plain <w:r> children and
    <w:r> children of <w:hyperlink>.
    """
    p_el = paragraph._p
    # Iterate direct children so we preserve document order across mixed
    # <w:r> and <w:hyperlink><w:r/></w:hyperlink> siblings.
    for child in p_el.iterchildren():
        tag = child.tag
        if tag == qn('w:r'):
            yield from _extract_from_run_element(child)
        elif tag == qn('w:hyperlink'):
            for r_el in child.iterchildren(qn('w:r')):
                yield from _extract_from_run_element(r_el)


def _extract_from_run_element(r_el):
    """Yield (text, bold) for each text node inside a raw <w:r> element."""
    # Bold detection: <w:rPr><w:b/></w:rPr> (or <w:b w:val="1"/>).
    bold = False
    rpr = r_el.find(qn('w:rPr'))
    if rpr is not None:
        b_el = rpr.find(qn('w:b'))
        if b_el is not None:
            val = b_el.get(qn('w:val'))
            bold = val in (None, '1', 'true')
    # <w:t>text</w:t> holds the actual visible text (there may be multiple).
    for t_el in r_el.iter(qn('w:t')):
        text = t_el.text or ''
        if text:
            yield text, bold

def _para_is_section(p):
    """Return True if this paragraph should start a new top-level section."""
    style = p.style.name if p.style else ''
    if 'Heading 1' in style:
        return True
    if style == 'Normal':
        # Use python-docx runs here (need font.size, which isn't cheap to grab
        # from raw XML). Section-title paragraphs never contain hyperlinks
        # so it's safe to skip the hyperlink walker for this check.
        runs = p.runs
        if runs and all((r.bold and r.font.size == HEADING1_FONT_SIZE_EMU) for r in runs if r.text.strip()):
            txt = p.text.strip()
            # Only promote short single-word/phrase titles; avoid promoting
            # long bold paragraphs that happen to use the same size.
            if txt and len(txt) <= 40:
                return True
    return False


def _para_is_subsection(p):
    style = p.style.name if p.style else ''
    return 'Heading 2' in style


def _para_is_list(p):
    style = p.style.name if p.style else ''
    return 'List Paragraph' in style


def _paragraph_to_obj(p):
    """Extract a docx paragraph to the schema dict.

    Concatenates run text WITHOUT separators (matches existing JSON format).
    Newlines inside runs are stripped so the definition-splitter on the client
    keeps working ('Term' bold run then 'Body' non-bold run become a defn).
    """
    text_parts = []
    runs = []
    for raw, bold in _iter_runs_including_hyperlinks(p):
        # Strip leading/trailing whitespace-newlines that arise from soft
        # line breaks; keep internal spaces intact.
        cleaned = raw.replace('\n', '').replace('\r', '')
        if not cleaned:
            continue
        runs.append({'text': cleaned, 'bold': bool(bold)})
        text_parts.append(cleaned)

    # Handle bullet-list paragraphs: prefix the text with a bullet char
    # so the client renderer recognizes it as a list item.
    is_list = _para_is_list(p)
    if is_list:
        joined = ''.join(text_parts).strip()
        if joined and not joined.startswith(('•', '·', '\u2022')):
            joined = '• ' + joined
            # Rewrite the first non-empty run to include the bullet
            if runs:
                runs[0] = {'text': '• ' + runs[0]['text'].lstrip(), 'bold': runs[0]['bold']}
        text = joined
    else:
        text = ''.join(text_parts)

    return {
        'text': text,
        'is_list': is_list,
        'runs': runs,
    }


def convert(docx_path, output_path=None):
    doc = Document(docx_path)
    sections = []
    current_section = None
    current_sub = None

    def _new_section(title):
        nonlocal current_section, current_sub
        current_section = {'title': title, 'intro': [], 'subs': []}
        current_sub = None
        sections.append(current_section)

    def _new_sub(title):
        nonlocal current_sub
        current_sub = {'title': title, 'paragraphs': []}
        current_section['subs'].append(current_sub)

    def _add_paragraph(p):
        obj = _paragraph_to_obj(p)
        if not obj['text'].strip() and not obj['runs']:
            return
        target = current_sub['paragraphs'] if current_sub else current_section['intro']
        target.append(obj)

    # Seed the preamble section (everything before first Heading 1).
    _new_section('_preamble_')

    # Skip the doc-title footer paragraph so it doesn't leak into JSON.
    SKIP_TEXTS = {'Version 3.1 · Updated July 28, 2026'}

    for p in doc.paragraphs:
        t = p.text.strip()

        if _para_is_section(p):
            _new_section(t)
            continue

        if _para_is_subsection(p):
            _new_sub(t)
            continue

        # Skip version-footer line entirely
        if t in SKIP_TEXTS:
            continue

        _add_paragraph(p)

    if output_path:
        Path(output_path).write_text(
            json.dumps(sections, indent=2, ensure_ascii=False),
            encoding='utf-8',
        )
    return sections


if __name__ == '__main__':
    docx = sys.argv[1] if len(sys.argv) > 1 else '/Users/jennamenking/Downloads/CW Product Methodologies July_28_2026A.docx'
    out = sys.argv[2] if len(sys.argv) > 2 else 'bg-webapp/static/cw_methodologies.json'
    sections = convert(docx, out)
    print(f'Wrote {out} with {len(sections)} sections:')
    for s in sections:
        subs = len(s.get('subs') or [])
        intro = len(s.get('intro') or [])
        print(f'  - {s["title"]!r}  (intro={intro}, subs={subs})')
